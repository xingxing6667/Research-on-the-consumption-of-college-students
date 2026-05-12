from __future__ import annotations

import re
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parent
SURVEY_DOCX = ROOT / "大学生消费情况调查问卷2－默认报告.docx"
OUT = ROOT / "output" / "extracted"


def parse_percent(text: str) -> float:
    if text is None:
        return np.nan
    text = str(text).strip().replace("%", "")
    if not text:
        return np.nan
    try:
        return float(text) / 100
    except ValueError:
        return np.nan


def extract_questionnaire(docx_path: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    questions: list[dict[str, str | int]] = []

    for idx, text in enumerate(paragraphs):
        if not re.match(r"^第\d+题", text):
            continue
        qtype = ""
        qtext = text
        qtype_match = re.search(r"\[(.*?)\]", text)
        if qtype_match:
            qtype = qtype_match.group(1)
            qtext = re.sub(r"\s*\[.*?\]\s*$", "", text)
        elif idx + 1 < len(paragraphs) and re.match(r"^\[.*?\]$", paragraphs[idx + 1]):
            qtype = paragraphs[idx + 1].strip("[]")

        m = re.match(r"^第(\d+)题\s*(.*)$", qtext)
        if not m:
            continue
        questions.append(
            {
                "question_id": int(m.group(1)),
                "question": m.group(2).strip(),
                "question_type": qtype,
            }
        )

    rows: list[dict[str, object]] = []
    for q, table in zip(questions, doc.tables):
        valid_n = np.nan
        parsed_rows = []
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if len(cells) < 3:
                continue
            option, count_text, percent_text = cells[:3]
            if "本题有效填写人次" in option:
                valid_n = pd.to_numeric(count_text, errors="coerce")
                continue
            parsed_rows.append((option, count_text, percent_text))

        for option, count_text, percent_text in parsed_rows:
            rows.append(
                {
                    **q,
                    "option": option,
                    "count": pd.to_numeric(count_text, errors="coerce"),
                    "percent": parse_percent(percent_text),
                    "percent_text": percent_text,
                    "valid_n": valid_n,
                }
            )

    return pd.DataFrame(rows), pd.DataFrame(questions)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    survey, questions = extract_questionnaire(SURVEY_DOCX)
    survey.to_csv(OUT / "questionnaire_summary.csv", index=False, encoding="utf-8-sig")
    questions.to_csv(OUT / "questionnaire_questions.csv", index=False, encoding="utf-8-sig")
    print(f"Extracted questionnaire data to: {OUT}")
    print(f"Questions: {len(questions)}, rows: {len(survey)}")


if __name__ == "__main__":
    main()
